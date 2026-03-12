import os
import io
import re

from google.auth.transport.requests import Request # pyright: ignore[reportMissingImports]
from google.oauth2.credentials import Credentials # pyright: ignore[reportMissingImports]
from googleapiclient.discovery import build # pyright: ignore[reportMissingImports]
from googleapiclient.http import MediaIoBaseUpload # pyright: ignore[reportMissingImports]
from google_auth_oauthlib.flow import Flow # pyright: ignore[reportMissingImports]
from docx import Document # pyright: ignore[reportMissingImports]
from docx.shared import Inches, Pt, RGBColor # pyright: ignore[reportMissingImports]
from docx.oxml import OxmlElement # pyright: ignore[reportMissingImports]
from docx.oxml.ns import qn # pyright: ignore[reportMissingImports]
from docx.oxml.shared import qn # pyright: ignore[reportMissingImports]
from docx.enum.text import WD_ALIGN_PARAGRAPH # pyright: ignore[reportMissingImports]
from googleapiclient.http import MediaIoBaseDownload # pyright: ignore[reportMissingImports]
from services.utils import get_sheet_column_index, Headers_Map, State_Keys_Map, SHEETS_RANGE_WIDTH

# Escopos de acesso necessários (Sheets, Docs e Drive)
SCOPES = [
    'https://www.googleapis.com/auth/spreadsheets',
    'https://www.googleapis.com/auth/documents',
    'https://www.googleapis.com/auth/drive'
]

# Deve ser exatamente a mesma URL que você colocou no Google Cloud Console
REDIRECT_URI = 'http://localhost:8501'

def authenticate_google():
    """Verifica se já existe token válido. Se não, retorna a URL para login."""
    creds = None
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
        
    if creds and creds.valid:
        return creds, None
        
    if creds and creds.expired and creds.refresh_token:
        creds.refresh(Request())
        with open('token.json', 'w') as token:
            token.write(creds.to_json())
        return creds, None
        
    # Se não tem credencial, gera o link para o botão do Streamlit
    flow = Flow.from_client_secrets_file('credentials.json', scopes=SCOPES, redirect_uri=REDIRECT_URI)
    auth_url, _ = flow.authorization_url(prompt='consent')
    
    return None, auth_url

def save_auth_token(code):
    """Pega o código que o Google devolve na URL e transforma no token.json."""
    flow = Flow.from_client_secrets_file('credentials.json', scopes=SCOPES, redirect_uri=REDIRECT_URI)
    flow.fetch_token(code=code)
    creds = flow.credentials
    with open('token.json', 'w') as token:
        token.write(creds.to_json())
    return creds

def get_services():
    """Inicializa e retorna os serviços do Sheets, Docs e Drive."""
    creds, _ = authenticate_google()
    sheets_service = build('sheets', 'v4', credentials=creds)
    docs_service = build('docs', 'v1', credentials=creds)
    drive_service = build('drive', 'v3', credentials=creds)
    return sheets_service, docs_service, drive_service

def create_drive_folder(drive_service, folder_name, parent_folder_id):
    """Cria uma nova pasta dentro de uma pasta específica no Google Drive."""
    file_metadata = {
        'name': folder_name,
        'mimeType': 'application/vnd.google-apps.folder',
        'parents': [parent_folder_id]
    }
    file = drive_service.files().create(body=file_metadata, fields='id').execute()
    return file.get('id')


def upload_image_to_drive(drive_service, file_stream, filename, folder_id=None):
    """Faz o upload de uma imagem para o Drive e a torna publicamente acessível para leitura."""
    file_metadata = {'name': filename}
    if folder_id:
        file_metadata['parents'] = [folder_id]
        
    media = MediaIoBaseUpload(file_stream, mimetype='image/jpeg', resumable=True)
    
    # 1. Faz o upload do arquivo
    file = drive_service.files().create(
        body=file_metadata,
        media_body=media,
        fields='id, webViewLink'
    ).execute()
    
    file_id = file.get('id')
    
    # Retorna o ID do arquivo e o link para visualização
    return file_id, file.get('webViewLink')

def check_client_in_sheets(sheets_service, spreadsheet_id, range_name, client):
    """
    Verifica se o cliente já existe.
    Retorna a linha onde ele está, o ID da pasta, o ID do doc, e os dados completos da linha (row).
    """
    sheet = sheets_service.spreadsheets()
    result = sheet.values().get(spreadsheetId=spreadsheet_id, range=range_name).execute()
    values = result.get('values', [])
    for i, row in enumerate(values):
        if len(row) > 0 and row[0] == client:
            while len(row) < SHEETS_RANGE_WIDTH:  # Garante que a row tenha o número correto de colunas, preenchendo com vazios se necessário
                row.append("")
            folder_id = row[Headers_Map.DRIVE_ID.column_index]
            diagnosis_doc_id = row[Headers_Map.DIAGNOSIS_DOC_ID.column_index]
            roadmap_doc_id = row[Headers_Map.ROADMAP_DOC_ID.column_index]
            return i + 1, folder_id, diagnosis_doc_id, roadmap_doc_id, row 
            
    # ADICIONADO UM NONE EXTRA AQUI PARA QUANDO NÃO ACHAR NADA:
    return -1, None, None, None, None

def save_to_sheets(sheets_service, spreadsheet_id, range_name, existing_row, row_data):
    """Salva os dados (Atualiza se a linha existir, ou adiciona uma nova)."""
    sheet = sheets_service.spreadsheets()
    body = {'values': [row_data]}
    
    if existing_row != -1:
        # ATUALIZA a linha existente
        # Extrai apenas a letra da coluna do range_name (ex: de 'Página1!A:Z' extrai 'A')
        coluna_inicial = range_name.split('!')[1].split(':')[0] if '!' in range_name else 'A'
        update_range = f"{range_name.split('!')[0]}!{coluna_inicial}{existing_row}"
        
        sheet.values().update(
            spreadsheetId=spreadsheet_id, range=update_range,
            valueInputOption="USER_ENTERED", body=body
        ).execute()
    else:
        # CRIA uma nova linha no final
        sheet.values().append(
            spreadsheetId=spreadsheet_id, range=range_name,
            valueInputOption="USER_ENTERED", body=body
        ).execute()

def get_headers_map(sheets_service, spreadsheet_id, range_name):
    """
    Retorna um dicionário onde a chave é o nome do cabeçalho 
    e o valor é o índice da coluna (0, 1, 2...).
    """
    try:
        # Extrai apenas o nome da aba do range (ex: 'DB Diagnóstico!A:AF' -> 'DB Diagnóstico')
        sheet_name = range_name.split('!')[0]
        
        # Busca apenas a primeira linha
        result = sheets_service.spreadsheets().values().get(
            spreadsheetId=spreadsheet_id, 
            range=f"'{sheet_name}'!1:1"
        ).execute()
        
        values = result.get('values', [])
        if not values:
            return {}
            
        headers = values[0]
        # Criamos o mapa: {'Nome do Cliente': 0, 'Data': 1, ...}
        return {header.strip(): i for i, header in enumerate(headers)}
    except Exception as e:
        print(f"Erro ao mapear cabeçalhos: {e}")
        return {}

def download_image_from_drive(drive_service, file_id):
    """Faz o download de um arquivo do Google Drive diretamente para a memória (BytesIO)."""
    request = drive_service.files().get_media(fileId=file_id)
    fh = io.BytesIO()
    downloader = MediaIoBaseDownload(fh, request)
    done = False
    while done is False:
        status, done = downloader.next_chunk()
    
    fh.seek(0) # Retorna o ponteiro para o início para o python-docx conseguir ler
    return fh

def set_cell_background(cell, fill_color):
    """Preenche o fundo de uma célula da tabela no Word (ex: '4D4D4D' para cinza escuro)."""
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def get_doc_content(docs_service, doc_id):
    """
    Recupera o conteúdo de texto de um documento do Google Docs.
    
    Args:
        doc_service: Instância do serviço da API (googleapiclient.discovery.build)
        doc_id: O ID único do documento no Google Drive.
        
    Returns:
        str: O texto completo do documento.
    """
    try:
        # Faz a chamada à API para recuperar o documento completo
        document = docs_service.documents().get(documentId=doc_id).execute()
        
        # O conteúdo principal fica dentro de 'body' -> 'content'
        doc_content = document.get('body').get('content')
        
        texto_extraido = ""
        
        # Itera sobre os elementos estruturais do documento
        for value in doc_content:
            # Verifica se o elemento é um parágrafo
            if 'paragraph' in value:
                elements = value.get('paragraph').get('elements')
                for elem in elements:
                    # Verifica se o elemento contém texto (textRun)
                    if 'textRun' in elem:
                        texto_extraido += elem.get('textRun').get('content')
            
            # Opcional: Você pode adicionar lógica para tabelas (table) aqui se necessário
            
        return texto_extraido

    except Exception as e:
        print(f"Ocorreu um erro ao ler o documento: {e}")
        return None

def add_doc_template(doc):
        # --- ADICIONAR HEADER ---
    try:
        # Acede à primeira secção do documento (que controla o cabeçalho)
        section = doc.sections[0]
        header = section.header

        # Pega no primeiro parágrafo do cabeçalho
        header_para = header.paragraphs[0]
        header_para.paragraph_format.space_before = Pt(0)
        header_para.paragraph_format.space_after = Pt(0)
        header_para.paragraph_format.left_indent = Pt(0)
        header_para.paragraph_format.right_indent = Pt(0)
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER # Centraliza a faixa
        
        # Insere a imagem da faixa no cabeçalho
        run = header_para.add_run()
        run.add_picture('images/image2.png', width=Inches(6))
        
    except Exception as e:
        print(f"Aviso: Não foi possível adicionar a imagem do cabeçalho: {e}")

    # --- ADICIONAR RODAPÉ ---
    try:
        footer = section.footer
        
        # Cria uma tabela de 1 linha e 2 colunas no rodapé. 
        # Por padrão no python-docx, essa tabela nasce sem bordas (invisível), o que é perfeito!
        footer_table = footer.add_table(rows=1, cols=2, width=Inches(6))
        
        # 1. Configura a imagem da Esquerda
        cell_left = footer_table.cell(0, 0)
        para_left = cell_left.paragraphs[0]
        para_left.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run_left = para_left.add_run()
        run_left.add_picture('images/image3.png', width=Inches(2.5)) 
        
        # 2. Configura a imagem da Direita
        cell_right = footer_table.cell(0, 1)
        para_right = cell_right.paragraphs[0]
        para_right.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run_right = para_right.add_run()
        run_right.add_picture('images/image1.png', width=Inches(2.5)) 
    except Exception as e:
        print(f"Aviso: Não foi possível adicionar as imagens do rodapé: {e}")

def create_update_diagnostic_doc(drive_service, client, diagnosis_data, folder_id, doc_id):
    """Cria/Atualiza o Google Docs estruturado usando python-docx, embutindo imagens nativamente."""

    def add_images(section, raw_images):      
        # Injeta as imagens cruas (UploadedFile do Streamlit) direto no Word!
        if raw_images:
            doc.add_paragraph("Evidências:")
            for img_file in raw_images:
                try:
                    img_file.seek(0) # Reinicia o ponteiro de leitura da imagem
                    image_stream = io.BytesIO(img_file.read())
                    doc.add_picture(image_stream, width=Inches(6.0)) # 6 polegadas ajusta bem na página
                    doc.add_paragraph()
                except Exception as e:
                    print(f"Aviso: Erro ao embutir imagem em {section}: {e}")
    
    # 1. Inicia o documento em memória
    doc = Document()

    add_doc_template(doc)
    
    # --- INTRODUÇÃO DO DOCUMENTO ---
    doc.add_heading(f"[DS30] Diagnóstico Técnico - {client}", level=1)

    doc.add_heading('Introdução', level=2)
    doc.add_paragraph('O projeto DS30, desenvolvido em parceria com o Google, visa a implementação e validação de soluções avançadas de Marketing Analytics. O foco central é o aprimoramento da mensuração e a otimização de dados por meio dos seguintes recursos:')
    
    itens_intro = ['Google Tag Gateway', 'Enhanced Conversions', 'Enhanced Conversions for Leads', 'Aprimoramento de integrações OCI', 'Google Analytics User-Provided Data (UPD)']
    for item in itens_intro:
        doc.add_paragraph(item, style='List Bullet')
        
    doc.add_paragraph('Este documento apresenta a etapa de diagnóstico, fundamental para identificar os requisitos técnicos e as ações necessárias para viabilizar a execução do projeto.')

    # --- SEÇÃO GOOGLE TAG GATEWAY ---
    doc.add_heading('Google Tag Gateway', level=2)
    doc.add_paragraph('Para a viabilização da implementação do Google Tag Gateway, é necessário entender o cenário de infraestrutura do cliente, onde será necessário adicionar um ou mais caminhos dedicados exclusivamente à mensuração e entender qual tecnologia é utilizada para a entrega de conteúdo web (CDN) é essencial para esta etapa.')
    doc.add_paragraph('Além disso, é necessário entender qual ferramenta Gerenciador de Tag é utilizada.')

    doc.add_heading('Diagnóstico', level=3)
    doc.add_paragraph('Através de formulários/entrevistas/reuniões e exploração do site, foi possível identificar que o cliente:')

    # Lógica do GTG (Mantendo exatamente a mesma estrutura)
    gtg_diagnosis_content = []
    gtg_diagnosis_content.append("Possui o Google Tag Gateway implementado;" if diagnosis_data.get(State_Keys_Map.GTG_IMPLEMENTED) == "Sim" else "Não possui o Google Tag Gateway implementado ou não informou;")
    gtg_diagnosis_content.append(f"Utiliza a CDN {diagnosis_data.get(State_Keys_Map.CDN)};" if diagnosis_data.get(State_Keys_Map.CDN) else "Cliente não utiliza ou não informou a CDN;")
    gtg_diagnosis_content.append("Utiliza IaC;" if diagnosis_data.get(State_Keys_Map.USE_IAC) == "Sim" else "Não utiliza IaC ou não informou;")
    gtg_diagnosis_content.append("Não utiliza TMS;" if diagnosis_data.get(State_Keys_Map.USE_TMS) != 'Sim' else (f'Utiliza TMS: {diagnosis_data.get(State_Keys_Map.TMS_TYPE)};' if diagnosis_data.get(State_Keys_Map.TMS_TYPE) != 'Selecione' else 'Utiliza TMS, mas não informou qual;'))
    
    if diagnosis_data.get(State_Keys_Map.TMS_TYPE) == "Google Tag Manager":
        gtg_diagnosis_content.append(f"IDs do GTM Client-Side: {diagnosis_data.get(State_Keys_Map.GTM_CS_IDS, 'Não informado')};")
        gtg_diagnosis_content.append("Não utiliza GTM Server-Side;" if diagnosis_data.get(State_Keys_Map.USE_GTM_SS) != 'Sim' else (f"Utiliza GTM Server-Side, provisionado no servidor {diagnosis_data.get(State_Keys_Map.GTM_SS_SERVER)};" if diagnosis_data.get(State_Keys_Map.GTM_SS_SERVER) else "Utiliza GTM Server-Side, mas não informou onde está provisionado;"))
        
    gtg_diagnosis_content.append(f"O cliente {'não ' if diagnosis_data.get(State_Keys_Map.AUTH_NEW_DOMAIN_PATH) != 'Sim' else ''}autorizou a criação de novos caminhos no domínio principal para a implementação do GTG.")

    for item in gtg_diagnosis_content:
        doc.add_paragraph(item, style='List Bullet')
    
    if diagnosis_data.get(State_Keys_Map.GTG_IMPLEMENTED) == 'Não':
        doc.add_heading("Considerações finais sobre o GTG", level=3)
        if diagnosis_data.get(State_Keys_Map.GTG_POSSIBLE_CONFIG) == "Sim":
            doc.add_paragraph(f"Foi possível concluir que será possível implementar o Google Tag Gateway no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.GTG_BLOCKS):
                doc.add_paragraph("Porém, foram identificados os seguintes bloqueios para a implementação do GTG:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.GTG_BLOCKS))
        elif diagnosis_data.get(State_Keys_Map.GTG_POSSIBLE_CONFIG) == "Não":
            doc.add_paragraph(f"Foi possível concluir que, à princípio, não será possível implementar o Google Tag Gateway no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.GTG_BLOCKS):
                doc.add_paragraph("Os bloqueios identificados para a implementação do GTG foram:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.GTG_BLOCKS))
        else:
            if diagnosis_data.get(State_Keys_Map.GTG_BLOCKS):
                doc.add_paragraph(f"Não foi possível concluir se será possível implementar o Google Tag Gateway no ambiente do cliente, pelos seguintes motivos a serem avaliados:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.GTG_BLOCKS))
            else:
                doc.add_paragraph(f"Não foi possível concluir se será possível implementar o Google Tag Gateway no ambiente do cliente.")

    if diagnosis_data.get(State_Keys_Map.GTG_PS):
        doc.add_heading("Observações - Google Tag Gateway", level=3)
        doc.add_paragraph(diagnosis_data.get(State_Keys_Map.GTG_PS))

    # --- SEÇÃO ENVIO DE SINAIS UPD ---

    doc.add_heading('Envio de Sinais UPD', level=2)
    doc.add_paragraph('A etapa de envio de sinais UPD tem como objetivo identificar oportunidades de implementação de User-Provided Data (UPD) para o Google Analytics, Enhanced Conversions e Enhanced Conversions for Leads.')
    doc.add_paragraph('O UPD é um recurso que permite o envio de dados adicionais sobre os usuários, enriquecendo a qualidade da mensuração e possibilitando uma melhor compreensão do comportamento do consumidor.')
    doc.add_paragraph('Nesta etapa, é necessário identificar se o cliente já possui algum tipo de implementação de UPD, seja por meio de hardcode ou por meio de ferramentas de Gerenciamento de Tags, e quais dados estão sendo enviados atualmente.')
    doc.add_paragraph('Além disso, é importante identificar oportunidades de implementação de UPD em formulários de contato, checkouts, ou outras interações relevantes no site.')
    doc.add_heading('Diagnóstico', level=3)
    doc.add_paragraph(f"{'Com a análise do arquivo JSON do container GTM, foi possível fazer as seguintes verificações:' if diagnosis_data.get(State_Keys_Map.GTM_ANALYSIS) else 'Não foi possível realizar a análise do container GTM, pois o cliente não forneceu o arquivo JSON exportado do container GTM Client-Side ou o cliente não utiliza GTM.'}")

    # Função auxiliar para criar tabelas estilizadas (Cabeçalho com fundo cinza escuro, letra branca e negrito)
    def add_styled_table(titulo, headers, data_rows):
        if not data_rows: return
        doc.add_paragraph(titulo)
        table = doc.add_table(rows=1, cols=len(headers))
        table.style = 'Table Grid'
        
        # Formata Cabeçalho (Fundo Cinza Escuro, Letra Branca, Negrito)
        hdr_cells = table.rows[0].cells
        for i, header in enumerate(headers):
            hdr_cells[i].text = header
            set_cell_background(hdr_cells[i], "4D4D4D")
            for paragraph in hdr_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.font.bold = True
                    run.font.color.rgb = RGBColor(255, 255, 255)
                    
        # Insere os dados
        for row_data in data_rows:
            row_cells = table.add_row().cells
            for i, cell_data in enumerate(row_data):
                row_cells[i].text = str(cell_data)
        doc.add_paragraph() # Dá um Enter abaixo da tabela

    if diagnosis_data.get(State_Keys_Map.GTM_ANALYSIS):
        upd_var_table_data = []
        target_tag_table_data = []
        for var in diagnosis_data.get(State_Keys_Map.GTM_ANALYSIS).get('upd_variables', {}).values():
            upd_var_table_data.append([var['var_name'], var['method']])
        for tag in diagnosis_data.get(State_Keys_Map.GTM_ANALYSIS).get('target_tags', {}).values():
            target_tag_table_data.append([tag['tag_name'], tag['type'], "⏸️ Sim" if tag['paused'] else "▶️ Não", tag['setting_status'], tag['var_name']])
        doc.add_heading("Análise do Container GTM", level=4)
        add_styled_table("Variáveis UPD Encontradas", ["Nome da Variável UPD", "Método"], upd_var_table_data)
        add_styled_table("Tags Relevantes Configuração UPD", ["Nome da Tag", "Tipo", "Pausada", "Status", "Variável UPD"], target_tag_table_data)

    if diagnosis_data.get(State_Keys_Map.FORM_URLS):
        doc.add_heading("Análise do site", level=4)
        doc.add_paragraph("Foi possível identificar as seguintes URLs de formulários relevantes para implementação de UPD:")
        for url in diagnosis_data.get(State_Keys_Map.FORM_URLS):
            doc.add_paragraph(url, style='List Bullet')

    # Função auxiliar para adicionar subseções de cada plataforma (GA, EC, ECL) mantendo a mesma estrutura e formatação
    def add_subsection(platform, status_plataforma):
        if status_plataforma:
            doc.add_paragraph(f"A configuração na interface da plataforma {platform}{'' if status_plataforma == 'Sim' else ' não'} está ativada corretamente.", style='List Bullet')
        else:
            doc.add_paragraph(f"Status da configuração na interface da plataforma {platform} não informado.", style='List Bullet')

    # --- SUBSEÇÃO GA UPD ---
    doc.add_heading("Google Analytics UPD", level=4)
    if diagnosis_data.get(State_Keys_Map.GA_HARDCODED):
        doc.add_paragraph("Foram identificados dados UPD hardcoded no código do site, o que pode indicar uma implementação manual de UPD para Google Analytics.", style='List Bullet')
    add_images("Envio de Sinais - Google Analytics", diagnosis_data.get(State_Keys_Map.GA_IMG_RAW))
    if diagnosis_data.get(State_Keys_Map.GA_UPD_CONFIG) == 'Não' and diagnosis_data.get(State_Keys_Map.GA_HARDCODED) == 'Não':
        doc.add_heading("Considerações finais sobre o GA UPD", level=5)
        if diagnosis_data.get(State_Keys_Map.GA_UPD_POSSIBLE_CONFIG) == "Sim":
            doc.add_paragraph(f"Foi possível concluir que será possível configurar o GA UPD no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.GA_UPD_BLOCKS):
                doc.add_paragraph("Porém, foram identificados os seguintes possíveis bloqueios para a configuração:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.GA_UPD_BLOCKS))
        elif diagnosis_data.get(State_Keys_Map.GA_UPD_POSSIBLE_CONFIG) == "Não":
            doc.add_paragraph(f"Foi possível concluir que, à princípio, não será possível configurar o GA UPD no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.GA_UPD_BLOCKS):
                doc.add_paragraph("Os bloqueios identificados para a configuração do GA UPD foram:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.GA_UPD_BLOCKS))
        else:
            if diagnosis_data.get(State_Keys_Map.GA_UPD_BLOCKS):
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o GA UPD no ambiente do cliente, pelos seguintes motivos a serem avaliados:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.GA_UPD_BLOCKS))
            else:
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o GA UPD no ambiente do cliente.")
    elif diagnosis_data.get(State_Keys_Map.GA_UPD_CONFIG) == 'Sim':
        doc.add_paragraph('A configuração na interface do GA foi executada corretamente.', style="List Bullet")

    
    # --- SUBSEÇÃO EC ---
    doc.add_heading("Enhanced Conversions", level=4)
    if diagnosis_data.get(State_Keys_Map.EC_HARDCODED):
        doc.add_paragraph("Foram identificados dados UPD hardcoded no código do site, o que pode indicar uma implementação manual de UPD para Enhanced Conversions.", style='List Bullet')
    if len(diagnosis_data.get(State_Keys_Map.EC_PLATFORMS, [])) > 0:
        for ec_platform in diagnosis_data.get(State_Keys_Map.EC_PLATFORMS, []):
            platform = ec_platform.split(" (")[0]  # Extrai o nome da plataforma, removendo o status entre parênteses
            platform_status = "Sim" if "Sim" in ec_platform else "Não" if "Não" in ec_platform else None
            add_subsection(platform, platform_status)
    else:
        doc.add_paragraph("O cliente não informou se possui alguma plataforma Google de marketing integrada que possa se beneficiar com o Enhanced Conversions.", style='List Bullet')
    add_images("Envio de Sinais - Enhanced Conversions", diagnosis_data.get(State_Keys_Map.EC_IMG_RAW))
    if "Não" in ", ".join(diagnosis_data.get(State_Keys_Map.EC_PLATFORMS, [])) and diagnosis_data.get(State_Keys_Map.EC_HARDCODED) == 'Não':
        doc.add_heading("Considerações finais sobre o EC", level=5)
        if diagnosis_data.get(State_Keys_Map.EC_POSSIBLE_CONFIG) == "Sim":
            doc.add_paragraph(f"Foi possível concluir que será possível configurar o EC no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.EC_BLOCKS):
                doc.add_paragraph("Porém, foram identificados os seguintes possíveis bloqueios para a configuração:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.EC_BLOCKS))
        elif diagnosis_data.get(State_Keys_Map.EC_POSSIBLE_CONFIG) == "Não":
            doc.add_paragraph(f"Foi possível concluir que, à princípio, não será possível configurar o EC no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.EC_BLOCKS):
                doc.add_paragraph("Os bloqueios identificados para a configuração do EC foram:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.EC_BLOCKS))
        else:
            if diagnosis_data.get(State_Keys_Map.EC_BLOCKS):
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o EC no ambiente do cliente, pelos seguintes motivos a serem avaliados:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.EC_BLOCKS))
            else:
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o EC no ambiente do cliente.")

    # --- SUBSEÇÃO ECL ---
    doc.add_heading("Enhanced Conversions for Leads", level=4)
    if diagnosis_data.get(State_Keys_Map.ECL_HARDCODED):
        doc.add_paragraph("Foram identificados dados UPD hardcoded no código do site, o que pode indicar uma implementação manual de UPD para Enhanced Conversions for Leads.", style='List Bullet')
    if len(diagnosis_data.get(State_Keys_Map.ECL_PLATFORMS, [])) > 0:
        for ecl_platform in diagnosis_data.get(State_Keys_Map.ECL_PLATFORMS, []):
            platform = ecl_platform.split(" (")[0]  # Extrai o nome da plataforma, removendo o status entre parênteses
            platform_status = "Sim" if "Sim" in ecl_platform else "Não" if "Não" in ecl_platform else None
            add_subsection(platform, platform_status)
    else:
        doc.add_paragraph("O cliente não informou se possui alguma plataforma Google de marketing integrada que possa se beneficiar com o Enhanced Conversions for Leads.", style='List Bullet')
    add_images("Envio de Sinais - Enhanced Conversions for Leads", diagnosis_data.get(State_Keys_Map.ECL_IMG_RAW))
    if "Não" in ", ".join(diagnosis_data.get(State_Keys_Map.ECL_PLATFORMS, [])) and diagnosis_data.get(State_Keys_Map.ECL_HARDCODED) == 'Não':
        doc.add_heading("Considerações finais sobre o ECL", level=5)
        if diagnosis_data.get(State_Keys_Map.ECL_POSSIBLE_CONFIG) == "Sim":
            doc.add_paragraph(f"Foi possível concluir que será possível configurar o ECL no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.ECL_BLOCKS):
                doc.add_paragraph("Porém, foram identificados os seguintes possíveis bloqueios para a configuração:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.ECL_BLOCKS))
        elif diagnosis_data.get(State_Keys_Map.ECL_POSSIBLE_CONFIG) == "Não":
            doc.add_paragraph(f"Foi possível concluir que, à princípio, não será possível configurar o ECL no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.ECL_BLOCKS):
                doc.add_paragraph("Os bloqueios identificados para a configuração do ECL foram:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.ECL_BLOCKS))
        else:
            if diagnosis_data.get(State_Keys_Map.ECL_BLOCKS):
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o ECL no ambiente do cliente, pelos seguintes motivos a serem avaliados:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.ECL_BLOCKS))
            else:
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o ECL no ambiente do cliente.")

    if diagnosis_data.get(State_Keys_Map.UPD_PS):
        doc.add_heading("Observações - Envio de Sinais UPD", level=3)
        doc.add_paragraph(diagnosis_data.get(State_Keys_Map.UPD_PS))

    # --- SEÇÃO OFFLINE CONVERSIONS INTEGRATION (OCI) ---
    doc.add_heading("Offline Conversions Integration (OCI)", level=2)
    doc.add_paragraph("A integração de Offline Conversions (OCI) é um processo que permite a importação de dados de conversões que ocorreram offline (como vendas em lojas físicas, chamadas telefônicas, etc.) para as plataformas de publicidade digital.")
    doc.add_paragraph("Isso é crucial para obter uma visão completa do desempenho das campanhas e otimizar as estratégias de marketing com base em dados mais abrangentes.")
    doc.add_paragraph("O diagnóstico nesta seção consiste em verificar se o cliente já possui alguma implementação de OCI, entender o método de integração utilizado (ex: API, upload manual, etc.), e quais informações estão sendo integradas atualmente.")
    doc.add_heading('Diagnóstico', level=3)
    oci_implementado = diagnosis_data.get(State_Keys_Map.OCI_IMPLEMENTED)
    doc.add_paragraph(f"O cliente {'possui' if oci_implementado == 'Sim' else 'não possui' if oci_implementado == 'Não' else 'não informou'} integração de Offline Conversions (OCI) implementada.", style='List Bullet')
    if oci_implementado == "Sim" and diagnosis_data.get(State_Keys_Map.OCI_METHOD):
            doc.add_paragraph(f"A integração de conversões offline é feita via {diagnosis_data.get(State_Keys_Map.OCI_METHOD)}.", style='List Bullet')
    if oci_implementado == "Sim" and diagnosis_data.get(State_Keys_Map.OCI_INFOS):
            dados = diagnosis_data.get(State_Keys_Map.OCI_INFOS).split("\n") if isinstance(diagnosis_data.get(State_Keys_Map.OCI_INFOS), str) else diagnosis_data.get(State_Keys_Map.OCI_INFOS)
            doc.add_paragraph(f"As seguintes informações estão sendo integradas atualmente:", style='List Bullet')
            for dado in dados:
                doc.add_paragraph(dado, style='List Bullet 2')
    add_images("Offline Conversions Integration (OCI)", diagnosis_data.get(State_Keys_Map.OCI_IMG_RAW))
    if diagnosis_data.get(State_Keys_Map.OCI_PS):
        doc.add_heading("Observações - Offline Conversion Integration (OCI)", level=3)
        doc.add_paragraph(diagnosis_data.get(State_Keys_Map.OCI_PS))
    if diagnosis_data.get(State_Keys_Map.OCI_IMPLEMENTED) == "Não":
        doc.add_heading("Considerações finais sobre o OCI", level=5)
        if diagnosis_data.get(State_Keys_Map.OCI_POSSIBLE_CONFIG) == "Sim":
            doc.add_paragraph(f"Foi possível concluir que será possível configurar o OCI no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.OCI_BLOCKS):
                doc.add_paragraph("Porém, foram identificados os seguintes possíveis bloqueios para a configuração:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.OCI_BLOCKS))
        elif diagnosis_data.get(State_Keys_Map.OCI_POSSIBLE_CONFIG) == "Não":
            doc.add_paragraph(f"Foi possível concluir que, à princípio, não será possível configurar o OCI no ambiente do cliente.")
            if diagnosis_data.get(State_Keys_Map.OCI_BLOCKS):
                doc.add_paragraph("Os bloqueios identificados para a configuração do OCI foram:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.OCI_BLOCKS))
        else:
            if diagnosis_data.get(State_Keys_Map.OCI_BLOCKS):
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o OCI no ambiente do cliente, pelos seguintes motivos a serem avaliados:")
                doc.add_paragraph(diagnosis_data.get(State_Keys_Map.OCI_BLOCKS))
            else:
                doc.add_paragraph(f"Não foi possível concluir se será possível configurar o OCI no ambiente do cliente.")

    # --- 2. SALVA NA MEMÓRIA ---
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    # --- 3. ENVIA PARA O GOOGLE DRIVE ---
    # O mimetype abaixo diz à API do Google: "Toma este ficheiro e transforma num Google Docs editável!"
    media = MediaIoBaseUpload(
        doc_bytes, 
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
        resumable=True
    )

    if doc_id:
        # Se o ficheiro já existir, este update SOBRESCREVE tudo, substituindo pelo conteúdo atualizado
        file = drive_service.files().update(
            fileId=doc_id,
            media_body=media
        ).execute()
    else:
        # Cria um ficheiro novo na pasta
        file_metadata = {
            'name': f'[DS30] Diagnóstico Técnico - {client}',
            'parents': [folder_id],
            'mimeType': 'application/vnd.google-apps.document'
        }
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        doc_id = file.get('id')

    # 4. Retorna os links formatados
    doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
    pdf_url = f"https://docs.google.com/document/d/{doc_id}/export?format=pdf"
    
    return doc_id, doc_url, pdf_url

def create_roadmap_doc(drive_service, client, folder_id, roadmap_data, doc_id):
    """Cria um documento Word isolado apenas com o Roadmap editado e envia para o Drive."""
    doc = Document()
    add_doc_template(doc)
    # --- CABEÇALHO (Opcional: Pode usar a mesma lógica de imagens se quiser) ---
    doc.add_heading(f"[DS30] Roadmap - {client}", level=1)
    p = doc.add_paragraph("Este documento apresenta o cronograma e o plano de ação sugerido com base no diagnóstico técnico das soluções de marketing analytics.")

    # --- FUNÇÃO AUXILIAR PARA INSERIR TEXTOS EDITADOS ---
    def inserir_sessao_roadmap(titulo, texto_editado):
        if texto_editado:
            doc.add_heading(titulo, level=2)
            for linha in texto_editado.split('\n'):
                if linha.strip():
                    content = re.sub(r'^\d+\.\s*', '', linha.strip(), flags=re.MULTILINE)
                    doc.add_paragraph(content, style='List Number')


    inserir_sessao_roadmap("Google Tag Gateway (GTG)", roadmap_data.get("gtg"))
    inserir_sessao_roadmap("Envio de Sinais (GA UPD / EC / ECL)", roadmap_data.get("upd"))
    inserir_sessao_roadmap("Integração de Conversões Offline (OCI)", roadmap_data.get("oci"))

    # --- SALVAR E FAZER UPLOAD ---
    doc_bytes = io.BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)

    media = MediaIoBaseUpload(
        doc_bytes, 
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document', 
        resumable=True
    )

    if doc_id:
        # Se o ficheiro já existir, este update SOBRESCREVE tudo, substituindo pelo conteúdo atualizado
        file = drive_service.files().update(
            fileId=doc_id,
            media_body=media
        ).execute()
    else:

        file_metadata = {
            'name': f'[DS30] Roadmap - {client}',
            'parents': [folder_id],
            'mimeType': 'application/vnd.google-apps.document'
        }
        file = drive_service.files().create(
            body=file_metadata,
            media_body=media,
            fields='id'
        ).execute()
        doc_id = file.get('id')
    doc_url = f"https://docs.google.com/document/d/{doc_id}/edit"
    pdf_url = f"https://docs.google.com/document/d/{doc_id}/export?format=pdf"
    
    return doc_id, doc_url, pdf_url